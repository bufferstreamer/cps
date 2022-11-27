import sys
from docx import Document
import pymysql
import datetime as dt
import os
import cpsSetting as c
from docx.shared import Pt
def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def replace_placeholder(doc, params):
    """替换占位符"""
    for paragraph in doc.paragraphs:
        for param in params:
            pv = str(params[param])
            ph = f'ph_{param}'
            if ph in paragraph.text:
                for run in paragraph.runs:
                    if ph in run.text:
                        run.text = run.text.replace(ph, pv)
                        run.italic = False


# 接收数据
path = sys.argv[1]
host = c.cps_host
user = c.cps_user
password = c.cps_password
# print(path)
# 获取当前绝对路径
docPath = os.getcwd()  # 获取当前路径
#docPath = 'E:\\IdeaProjects\\cps-3'
# print(docPath)
# docPath.replace("otherPython","tender")
print("start")
print(docPath)
doc = Document(docPath + "/otherPython/招标文件模板.docx")
# doc = Document("D:\idea\project\cps\document\\tenderC\招标文件模板.docx")
print("connect first database")
db_cps = pymysql.connect(host=host,
                         user=user,
                         password=password,
                         database='cps')
print("connect second database")
db_fmmall = pymysql.connect(host=host,
                         user=user,
                         password=password,
                         database='fmmall2')
cursor_cps = db_cps.cursor()
cursor_fmmall = db_fmmall.cursor()
print("connect finish")
sql_select_tender = "select id,\
                        tender_name,\
                        text,\
                        receive_time,\
                        deliver_point,\
                        project_text,\
                        license,\
                        finance,\
                        illegal,\
                        reputation_requirements,\
                        bidder_text,\
                        documents_obtain,\
                        bid_ddl,\
                        bid_submission,\
                        other_platform,\
                        electronic_bidding_rules,\
                        initiator,\
                        contacts,\
                        contact_number,\
                        email,\
                        address,\
                        qualification_review_time,\
                        bidding_start_time,\
                        bidding_end_time,\
                        publication_time,\
                        create_time\
                    from trender_information\
                    where id = %s" % (path)
cursor_cps.execute(sql_select_tender)
result_tender = cursor_cps.fetchall()
params = {
    "tender_name": result_tender[0][1],  # "一号",
    "text": result_tender[0][2],  # "用于测试招标文件生成",
    "receive_time": result_tender[0][3].date(),  # "2022年10月23日",
    "deliver_point": result_tender[0][4],  # "送达地点",
    "project_text": result_tender[0][5],  # "无",
    "license": result_tender[0][6],  # "需要营业执照",
    "finance": result_tender[0][7],  # "需要规模在500万以上",
    "illegal": result_tender[0][8],  # "需无犯罪记录",
    "reputation_requirements": result_tender[0][9],  # "信誉分数在50以上",
    "bidder_text": result_tender[0][10],  # "无",
    "documents_obtain": result_tender[0][11],  # "在各招标平台获取",
    "bid_ddl": result_tender[0][12].date(),  # "2022年10月25日",
    "bid_submission": result_tender[0][13],  # "综合楼409",
    "other_platform": result_tender[0][14],  # "腾讯招标、支付宝招标",
    "electronic_bidding_rules": result_tender[0][15],  # "根据说明书完成招标工作",
    "initiator": result_tender[0][16],  # "admin",
    "contacts": result_tender[0][17],  # "张三",
    "contact_number": result_tender[0][18],  # "1888888888",
    "email": result_tender[0][19],  # "154554545@111.com",
    "address": result_tender[0][20],  # "辽宁大连",
    "qualification_review_time": result_tender[0][21].date(),  # "2022年10月26日",
    "bidding_start_time": result_tender[0][22].date(),  # "2022年10月27日",
    "bidding_end_time": result_tender[0][23].date(),  # "2022年10月27日",
    "publication_time": result_tender[0][24].date(),  # "2022年10月27日",
    "create_time": result_tender[0][25].date(),  # "2022年10月01日",
}

replace_placeholder(doc, params)

table_title = ["产品id", '产品名称', '商品规格名称','商品属性','指标名称', '指标要求', '需求数量', '商品单位','预算金额']  # 定义表格的第一行的标题
table = doc.add_table(rows=1, cols=9)  # 定义表格的行数、列数
table.autofit =True
table.style = doc.styles["Table Grid"]
table_cells = table.rows[0].cells  # 将 table_title 的每列的名称写入表格
table_cells[0].text = table_title[0]
table_cells[1].text = table_title[1]
table_cells[2].text = table_title[2]
table_cells[3].text = table_title[3]
table_cells[4].text = table_title[4]
table_cells[5].text = table_title[5]
table_cells[6].text = table_title[6]
table_cells[7].text = table_title[7]
table_cells[8].text = table_title[8]
for i in table.rows[0].cells:
    run = i.paragraphs[0].runs[0]
    run.font.size = Pt(8)
# run = table.rows[0].cells.
# run.font.size = 10

sql_select_seed = "select product_id,\
                          product_name,\
                          sku_name,\
                          untitled,\
                          shop_index_name,\
                          shop_index_require,\
                          plan_number,\
                          goods_unit,\
                          plan_price\
                    from tender_seed\
                    where tender_information_id = %s" % (path)
cursor_fmmall.execute(sql_select_seed)
data = list(cursor_fmmall.fetchall())

# data = [            # 定义 data 的内容，准备将其追加写入表格
# ('1','红烧牛肉方便面', '包装规格', '五连袋','100件','2000元'),
# ('2','依云矿泉水', '容量', '500ml','300瓶','3000元'),
# ('3','养乐多', '包装规格', '五连装','50件','600元')
# ]

for i in data:  # 利用 for 循环将 data 追加写入表格
    row_cells = table.add_row().cells
    row_cells[0].text = str(i[0])
    row_cells[1].text = i[1]
    row_cells[2].text = i[2]
    row_cells[3].text = i[3]
    row_cells[4].text = i[4]
    row_cells[5].text = i[5]
    row_cells[6].text = str(i[6])
    row_cells[7].text = i[7]
    row_cells[8].text = str(i[8])

expect_text = '所需商品及商品需求如下：'

for paragraph in doc.paragraphs:

    paragraph_text = paragraph.text

    if paragraph_text.endswith(expect_text):
        target = paragraph
        break

move_table_after(table, target)
print("saving")
doc.save(docPath + "/otherPython/biddingDocuments/" + path + result_tender[0][1] + "招标文件.docx")
print("bidding create finish")
db_cps.close()
db_fmmall.close()
#doc.save(docPath + "/cps-admin/src/main/resources/static/documentTemplate/tender/" + path + result_tender[0][1] + "招标文件.docx")
# cps-admin/src/main/resources/static/documentTemplate/tender
# doc.save("D:\idea\project\cps\document/tender/" + result_tender[0][1] + "招标文件.docx")
